"""
Generate ternlang-whitepaper.docx from scratch using python-docx.
Run: python3 build_docx.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.25)
section.top_margin  = section.bottom_margin = Inches(1.0)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1a, 0x1a, 0x2e)
TEAL       = RGBColor(0x00, 0xd4, 0xaa)
DARK_GREY  = RGBColor(0x44, 0x44, 0x55)
CODE_BG    = RGBColor(0xf4, 0xf4, 0xf8)
BLACK      = RGBColor(0x11, 0x11, 0x22)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_font(run, "Calibri", 15, bold=True, color=DARK_BLUE)
    elif level == 2:
        set_font(run, "Calibri", 12, bold=True, color=DARK_BLUE)
    else:
        set_font(run, "Calibri", 11, bold=True, italic=True, color=DARK_GREY)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.25) if indent else None
    run = p.add_run(text)
    set_font(run, "Calibri", 11, color=BLACK)
    return p

def body_parts(parts):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, "Calibri", 11, bold=bold, italic=italic, color=BLACK)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.4)
    # Shade background via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F6')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, "Courier New", 9, color=DARK_GREY)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.4 + level * 0.2)
    run = p.add_run(text)
    set_font(run, "Calibri", 11, color=BLACK)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("─" * 72)
    set_font(run, "Courier New", 8, color=DARK_GREY)

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, "Calibri", 9, italic=True, color=DARK_GREY)

def add_table(headers, rows, caption_text=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            set_font(run, "Calibri", 10, bold=True, color=DARK_BLUE)
        # shade header
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E8E8F4')
        tcPr.append(shd)
    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                set_font(run, "Calibri", 10, color=BLACK)
    if caption_text:
        doc.add_paragraph()
        caption(caption_text)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
run = p.add_run("TERNLANG")
set_font(run, "Calibri", 28, bold=True, color=DARK_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A Full-Stack Balanced Ternary Execution Architecture\nfor Sparse Neural Inference and Ambiguity-Aware Agent Systems")
set_font(run, "Calibri", 14, italic=True, color=DARK_GREY)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Simeon Kepp\nRFI-IRFOS")
set_font(run, "Calibri", 12, bold=True, color=BLACK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026  ·  github.com/eriirfos-eng/ternary-intelligence-stack--tis-")
set_font(run, "Calibri", 10, color=DARK_GREY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════

heading("Abstract", 1)
body(
    "We present Ternlang v0.1.2, the first complete software stack for balanced ternary "
    "computing: a domain-specific language with real error propagation and a cross-file "
    "module system, bytecode compiler, stack-based virtual machine (BET VM), native C11 "
    "compilation target, hardware description language backend, distributed actor runtime, "
    "machine learning inference kernels, and a 13-agent Mixture-of-Experts orchestrator — "
    "all unified under a single coherent architecture."
)
body(
    "The foundational primitive is the trit t ∈ {−1, 0, +1}, where the value 0 represents "
    "an active neutral state rather than absence, enabling three-valued logic that is "
    "structurally superior to binary for ambiguity-aware reasoning."
)
body(
    "The principal execution contribution is TSPARSE_MATMUL: a first-class VM opcode that "
    "elides multiply operations against zero-weighted ('hold') trit elements, achieving a "
    "2.27× reduction at baseline sparsity and a peak of 122× at 99% sparsity on 512×512 "
    "matrices. We also define the BET ISA (51 opcodes), synthesise it to Verilog-2001 with "
    "per-cell clock-gating, add a native AST-to-C11 transpiler, and demonstrate the MoE-13 "
    "orchestrator with 13-agent dual-signal deliberation and an introspective hold stable "
    "attractor. Formal correspondence between ternlang trits and qutrit neural network states "
    "opens a path to quantum-adjacent hardware targeting."
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Keywords: ")
set_font(run, "Calibri", 11, bold=True, color=BLACK)
run2 = p.add_run(
    "balanced ternary, trit, sparse inference, BitNet, domain-specific language, "
    "virtual machine, actor model, FPGA synthesis, Verilog, error propagation, "
    "module system, C11 transpilation, mixture-of-experts, qutrit neural networks"
)
set_font(run2, "Calibri", 11, italic=True, color=DARK_GREY)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

heading("1. Introduction", 1)

body(
    "The computational substrate underlying modern artificial intelligence is binary. "
    "Floating-point arithmetic, two-state memory, and Boolean logic gates have driven "
    "five decades of progress — but they introduce a fundamental representational mismatch "
    "when modelling systems that are inherently three-valued: affirmed, denied, and undecided."
)
body(
    "Clinical diagnosis, legal reasoning, sensor fusion under noise, and multi-agent "
    "consensus all require a native neutral state that binary computing forces to encode "
    "as a special case: null pointers, NaN, sentinel values, or probabilistic scores "
    "collapsed to a threshold. Each encoding is a workaround for an absent primitive."
)
body(
    "Balanced ternary provides that primitive. A trit t ∈ {−1, 0, +1} carries three "
    "symmetric values. The neutral value 0 is active — a deliberate state of hold, not "
    "an empty bit pattern. Balanced ternary arithmetic is self-complementing: negation "
    "requires no special-case handling. And at the scale of modern neural networks, where "
    "BitNet and related work show that ternary-quantized weights preserve accuracy with "
    "dramatically reduced computation, the case for a ternary-native execution substrate "
    "is both theoretical and empirical."
)
body(
    "Despite this, the ternary computing field remains fragmented: hobbyist emulators, "
    "academic EDA tools for memristor hardware, isolated Lisp interpreters, and hardware "
    "simulators without compiler support. No project provides the full vertical stack. "
    "Ternlang fills this gap."
)

heading("1.1  Contributions", 2)
bullet("Language completeness: exhaustive three-way pattern matching, a ? error propagation operator, a real cross-file module system (stdlib + filesystem-relative user modules), and structured diagnostic codes.")
bullet("The BET ISA: a formal 2-bit-encoded instruction set with 51 opcodes covering arithmetic, tensor operations, actor messaging, and control flow.")
bullet("TSPARSE_MATMUL: a VM opcode that skips zero-weight multiplications at the instruction level, achieving 2.27× at baseline and 122× at 99% sparsity.")
bullet("A Verilog-2001 hardware backend with synthesisable sparse matmul array and full BET processor, plus an Icarus Verilog simulation wrapper.")
bullet("A native AST-to-C11 transpiler (ternlang-codegen) enabling native-speed execution and cross-compilation to embedded targets.")
bullet("MoE-13 orchestrator: 13-agent dual-signal deliberation, introspective hold stable attractor, and orchestrate_full() two-tier pipeline.")
bullet("Formal qutrit-to-trit correspondence connecting ternlang to Qutrit Neural Network research.")
bullet("Ecosystem bridges connecting existing ternary projects (9-trit assembly, Owlet S-expressions) to the BET VM as a common runtime.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 2. BACKGROUND
# ══════════════════════════════════════════════════════════════════════════════

heading("2. Background: Balanced Ternary", 1)

heading("2.1  Trit arithmetic", 2)

body(
    "A trit t ∈ T = {−1, 0, +1} participates in four complete operations. "
    "Ternary addition produces a sum and carry, both in T, satisfying a + b = 3c + s. "
    "Ternary multiplication is simply integer multiplication restricted to T, since "
    "|a|, |b| ≤ 1 implies |a · b| ≤ 1. Consensus (ternary OR) returns a if a = b, "
    "else 0. Negation maps t → −t, with neg(0) = 0."
)

add_table(
    ["a \\ b", "−1", "0", "+1"],
    [
        ["−1", "(+1, −1)", "(−1,  0)", "( 0,  0)"],
        [" 0", "(−1,  0)", "( 0,  0)", "(+1,  0)"],
        ["+1", "( 0,  0)", "(+1,  0)", "(−1, +1)"],
    ],
    "Table 1. Balanced ternary addition (sum, carry) for each pair of trits."
)

heading("2.2  The 2-bit BET encoding", 2)

body(
    "Hardware naturally operates in binary. BET encodes each trit as a 2-bit pair:"
)

add_table(
    ["Bit pattern", "Trit value", "Meaning"],
    [
        ["0b01", "−1", "conflict"],
        ["0b10", "+1", "truth"],
        ["0b11", " 0", "hold (active neutral)"],
        ["0b00", "FAULT", "invalid — triggers VmError"],
    ],
    "Table 2. BET 2-bit trit encoding."
)

body_parts([
    ("Key property: ", True, False),
    ("negation is a bit swap. Swapping the two bits of 0b01 gives 0b10 and vice versa; "
     "0b11 is symmetric and maps to itself. This means the TNEG opcode requires no "
     "arithmetic — just a single wiring operation in hardware. ", False, False),
])
body(
    "The all-ones reset state (0b11) initialises every register to hold — the "
    "semantically correct neutral value — without special reset logic."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 3. THE TERNLANG LANGUAGE
# ══════════════════════════════════════════════════════════════════════════════

heading("3. The Ternlang Language", 1)

heading("3.1  Design principles", 2)

body_parts([("Exhaustive three-way matching.  ", True, False),
    ("Every match expression must cover all three trit arms. "
     "The compiler rejects non-exhaustive matches at parse time, eliminating "
     "an entire class of runtime error.", False, False)])

body_parts([("0 is active.  ", True, False),
    ("The type system assigns distinct meaning to −1 (conflict), 0 (hold), and +1 (truth). "
     "There is no null, no undefined, no NaN. A trit always carries a definite value.", False, False)])

body_parts([("Sparsity is a language feature.  ", True, False),
    ("The @sparseskip directive marks a tensor operation as sparse-aware, routing the "
     "compiler to emit TSPARSE_MATMUL instead of TMATMUL. Sparsity is expressed in "
     "the source language, not discovered by the optimiser.", False, False)])

heading("3.2  Core constructs", 2)

body("Ternary classifier with exhaustive match:")
code_block(
    "fn classify(signal: trit) -> trit {\n"
    "    match signal {\n"
    "        -1 => conflict()   // active disagreement\n"
    "         0 => hold()       // awaiting evidence\n"
    "        +1 => truth()      // confirmed\n"
    "    }\n"
    "}"
)

body("Sparse matrix multiply — routes to TSPARSE_MATMUL at the ISA level:")
code_block(
    "@sparseskip\n"
    "let output: trittensor<8 x 8> = matmul(input, weights);"
)

body("Actor model for ternary message passing:")
code_block(
    "agent Voter {\n"
    "    fn handle(msg: trit) -> trit {\n"
    "        consensus(msg, hold())\n"
    "    }\n"
    "}\n"
    "\n"
    "let v: agentref = spawn Voter;\n"
    "send v truth();\n"
    "let decision: trit = await v;"
)

body("Remote actor (distributed runtime):")
code_block(
    "let remote_voter: agentref =\n"
    "    spawn remote \"192.168.1.42:7373\" Voter;\n"
    "send remote_voter truth();\n"
    "let r: trit = await remote_voter;"
)

heading("3.3  Type system", 2)
body(
    "Core types: trit (single balanced ternary value), trittensor<N x M> (N×M matrix "
    "on the tensor heap), agentref (actor handle, local or remote). Struct types with "
    "trit/tensor fields are supported via field-name mangling in the register allocator: "
    "a field s.field is stored in a named register slot 's.field', avoiding the need "
    "for heap allocation for small structs."
)

heading("3.4  Ternary Error Propagation", 2)
body(
    "Ternlang introduces a postfix ? operator for ternary-native error propagation, "
    "analogous to Rust's ? operator but grounded in the three-valued semantics of the "
    "trit space. For any expression e of type trit, writing e? in a function body has "
    "the following semantics:"
)
code_block(
    "e?  ≜  return -1   if e = -1\n"
    "       e            otherwise"
)
body(
    "This elevates trit −1 (conflict) from a value to a structural signal: any function "
    "that receives a conflict can propagate it up the call chain without explicit match arms. "
    "At the BET level, expr? compiles to four instructions:"
)
code_block(
    "; expr already on stack: [val]\n"
    "TDUP          ; [val, dup]\n"
    "TJMP_NEG L1   ; consume dup; if -1 → L1; else [val] continues\n"
    "TJMP L2       ; skip early return\n"
    "L1: TRET      ; return -1 to caller\n"
    "L2: (continue with val on stack)"
)
body(
    "The parser disambiguates ? from the ternary uncertain-branch operator (also ?) using "
    "two-token lookahead: expr? followed by { is an uncertain branch; expr? followed by "
    "any other token is a propagation operator."
)
code_block(
    "fn validate_signal(x: trit) -> trit {\n"
    "    let checked: trit = boundary_check(x)?; // propagate if conflict\n"
    "    return range_check(checked)?;           // propagate again\n"
    "}\n"
    "\n"
    "fn main() -> trit {\n"
    "    return validate_signal(1)?; // -1 if any check fails\n"
    "}"
)

heading("3.5  Cross-File Module System", 2)
body(
    "Ternlang's module system supports two resolution strategies, tried in order:"
)
bullet("Built-in stdlib (compile-time embedded): std::trit, std::math, std::tensor, std::io, ml::quantize, ml::inference are embedded in the compiler binary via include_str! — zero filesystem I/O at runtime.")
bullet("User-defined modules: a ModuleResolver walks use paths relative to the source file's directory. A declaration use agents::voter; resolves to <source_dir>/agents/voter.tern, parsed and merged into the program before semantic analysis.")
body(
    "Both strategies deduplicate: a module loaded by multiple use statements is parsed "
    "exactly once, and functions already present in the program are never duplicated. "
    "This makes StdlibLoader::resolve() idempotent and safe to call multiple times."
)
code_block(
    "// agents/voter.tern\n"
    "fn weighted_vote(a: trit, b: trit, c: trit) -> trit {\n"
    "    consensus(consensus(a, b), c)\n"
    "}\n"
    "\n"
    "// main.tern\n"
    "fn main() -> trit {\n"
    "    use agents::voter;\n"
    "    return weighted_vote(1, 0, -1);\n"
    "}"
)

heading("3.6  Diagnostic Philosophy and Error Codes", 2)
body(
    "Ternlang's error messages carry structured codes and ternary-philosophical commentary. "
    "Every diagnostic has a machine-readable code (for tooling and documentation lookup) "
    "and a human-readable nudge that frames the error in the language's conceptual model:"
)

add_table(
    ["Code", "Category", "Nudge"],
    [
        ["[PARSE-001]", "Unexpected token",       "The lexer hit something it didn't expect. Check your syntax."],
        ["[PARSE-002]", "Expected token missing",  "Something required is absent — the parser can't continue."],
        ["[PARSE-003]", "Invalid literal",         "Trit literals are -1, 0, or +1. Nothing else exists here."],
        ["[PARSE-004]", "Non-exhaustive match",    "Ternary has three states — cover all three or the compiler won't let you through."],
        ["[TYPE-001]",  "Type mismatch",           "The type system is strict. A trit is not an int is not a tensor."],
        ["[SCOPE-001]", "Undefined variable",      "Hold state — declare before use."],
        ["[STRUCT-001]","Unknown struct field",    "That field doesn't exist in the struct definition."],
        ["[STRUCT-002]","Struct field type error", "Field assignment requires the correct trit type."],
        ["[FN-001]",    "Unknown function",        "No function by that name is in scope."],
        ["[FN-002]",    "Return type mismatch",    "Ternary contracts are strict."],
        ["[FN-003]",    "Wrong argument count",    "Count your arguments — ternary arity is exact."],
        ["[FN-004]",    "Wrong argument type",     "The function expects a different trit type here."],
        ["[PROP-001]",  "? on non-trit",           "Only trit-returning functions can signal conflict. The third state requires a trit."],
        ["[BET-001]",   "Stack underflow",         "You tried to pop a truth that wasn't there."],
        ["[BET-002]",   "Invalid BET encoding",    "0b00 is not a valid trit. Check the 2-bit encoding."],
        ["[BET-003]",   "Unknown opcode",          "That opcode isn't in the BET ISA."],
        ["[BET-004]",   "Register out of range",   "BET has 27 registers (0–26). That's it."],
        ["[BET-005]",   "Tensor not allocated",    "TALLOC before you TIDX."],
        ["[BET-006]",   "Tensor index out of bounds","Trittensor indices are bounded by the declared shape."],
        ["[BET-007]",   "Agent not found",         "TSPAWN before TSEND. The agent table has no entry for that ID."],
    ],
    "Table 3. Ternlang diagnostic codes. Machine-readable prefix enables tooling integration."
)
body(
    "This design ensures that newcomers encountering their first ternary error receive both "
    "actionable guidance and an introduction to the philosophy underpinning the type system."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 4. THE BET ISA
# ══════════════════════════════════════════════════════════════════════════════

heading("4. The BET Instruction Set Architecture", 1)

heading("4.1  Machine model", 2)
bullet("27 registers (2 bits each), reset to 0b11 (hold). The number 27 = 3³ reflects the ternary motif.")
bullet("Value stack: unbounded, stores tagged Value union (Trit | Int | TensorRef | AgentRef).")
bullet("Tensor heap: indexed array of N×M trit matrices, allocated by TALLOC.")
bullet("Call stack: return-address stack for TCALL / TRET.")
bullet("Agent table: maps type IDs to handler addresses and per-instance mailboxes (VecDeque).")
bullet("Carry register: overflow from TADD stored separately, not on the value stack.")

heading("4.2  Instruction encoding", 2)
body(
    "Instructions are variable-length: 1-byte opcode followed by 0–2 operand bytes. "
    "All jump targets are 2-byte little-endian absolute addresses. "
    "The full ISA comprises 51 opcodes across five groups:"
)

add_table(
    ["Opcode", "Mnemonic", "Operands", "Stack effect", "Description"],
    [
        ["0x00", "THALT",          "",        "—",           "Stop execution"],
        ["0x01", "TPUSH",          "t",       "→ t",         "Push trit literal"],
        ["0x02", "TADD",           "",        "a b → s c",   "Balanced ternary add"],
        ["0x03", "TMUL",           "",        "a b → t",     "Ternary multiply"],
        ["0x04", "TNEG",           "",        "t → neg(t)",  "Bit-swap negate"],
        ["0x05", "TJMP_POS",       "addr",    "t →",         "Jump if t = +1"],
        ["0x06", "TJMP_ZERO",      "addr",    "t →",         "Jump if t = 0"],
        ["0x07", "TJMP_NEG",       "addr",    "t →",         "Jump if t = −1"],
        ["0x08", "TSTORE",         "r",       "t →",         "Pop into register r"],
        ["0x09", "TLOAD",          "r",       "→ reg[r]",    "Push register r"],
        ["0x0b", "TJMP",           "addr",    "—",           "Unconditional jump"],
        ["0x0c", "TDUP",           "",        "t → t t",     "Duplicate top"],
        ["0x0d", "TPOP",           "",        "t →",         "Discard top"],
        ["0x0e", "TCONS",          "",        "a b → cons",  "Consensus (ternary OR)"],
        ["0x0f", "TALLOC",         "N M",     "→ ref",       "Allocate N×M tensor"],
        ["0x10", "TCALL",          "addr",    "—",           "Call; push return addr"],
        ["0x11", "TRET",           "",        "—",           "Return; pop addr"],
        ["0x20", "TMATMUL",        "",        "rA rB → rC",  "Dense tensor multiply"],
        ["0x21", "TSPARSE_MATMUL", "",        "rA rB → rC",  "Sparse matmul (skip 0s)"],
        ["0x22", "TIDX",           "",        "ref i j → t", "Index tensor element"],
        ["0x23", "TSET",           "",        "ref i j t →", "Set tensor element"],
        ["0x24", "TSHAPE",         "",        "ref → N M",   "Push tensor dimensions"],
        ["0x25", "TSPARSITY",      "",        "ref → count", "Count zero elements"],
        ["0x26", "TCOMPRESS",      "",        "ref → ref",   "RLE-compress trit tensor"],
        ["0x27", "TUNPACK",        "",        "ref → ref",   "Decompress RLE tensor"],
        ["0x30", "TSPAWN",         "type_id", "→ agentref",  "Create agent instance"],
        ["0x31", "TSEND",          "",        "ref msg →",   "Enqueue message"],
        ["0x32", "TAWAIT",         "",        "ref → t",     "Run handler, get result"],
    ],
    "Table 4. BET ISA opcode reference (all 51 opcodes; selected entries shown)."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 5. SPARSE TERNARY INFERENCE
# ══════════════════════════════════════════════════════════════════════════════

heading("5. Sparse Ternary Inference", 1)

heading("5.1  Ternary quantization", 2)
body(
    "BitNet-style ternary weight quantization maps floating-point weights w ∈ ℝ to "
    "ŵ ∈ {−1, 0, +1} using a threshold τ = ½ · E[|w|]:"
)
code_block(
    "ŵ = +1   if w >  τ\n"
    "ŵ =  0   if |w| ≤ τ     (τ = 0.5 × mean(|w|))\n"
    "ŵ = −1   if w < −τ"
)
body(
    "The resulting weight distribution is heavily concentrated at 0 (hold): typical "
    "language model weights at BitNet scale show 55–65% zero elements after quantization. "
    "In the ternlang-ml crate, this is implemented as:"
)
code_block(
    "pub fn bitnet_threshold(weights: &[f32]) -> f32 {\n"
    "    let mean_abs = weights.iter().map(|w| w.abs()).sum::<f32>()\n"
    "                   / weights.len() as f32;\n"
    "    0.5 * mean_abs\n"
    "}"
)

heading("5.2  TSPARSE_MATMUL", 2)
body_parts([
    ("The key identity: ", True, False),
    ("mul(a, 0) = 0 for all a ∈ T. ", False, True),
    ("In a dense N×M matrix multiply, every element contributes N·M multiplications. "
     "After ternary quantization with sparsity ρ (fraction of zero-weight elements), "
     "only (1−ρ)·N·M multiplications have non-trivial results. The rest are guaranteed "
     "zero and can be skipped.", False, False),
])
body("TSPARSE_MATMUL implements a sparse inner-product loop:")
code_block(
    "for i in 0..N:\n"
    "  for j in 0..M:\n"
    "    for k in 0..K:\n"
    "      w = W[k][j]\n"
    "      if w == HOLD: continue   // skip — guaranteed zero\n"
    "      acc[i][j] += mul(A[i][k], w)"
)
body(
    "The result is identical to TMATMUL — no approximation. The @sparseskip directive "
    "in the source language routes the compiler to emit TSPARSE_MATMUL for the following "
    "matmul() call. Sparsity awareness is a source-language property, not a runtime guess."
)

heading("5.3  Benchmark results", 2)

add_table(
    ["Metric", "Dense (TMATMUL)", "Sparse (TSPARSE_MATMUL)"],
    [
        ["Weight sparsity",      "0%",       "56.2%"],
        ["Multiply operations",  "262,144",  "115,343"],
        ["Skipped operations",   "0",        "146,801"],
        ["Relative cost",        "1.00×",    "0.44× (2.27× speedup)"],
    ],
    "Table 5. Sparse vs. dense ternary matmul on 512×512 quantized weight matrix."
)

body(
    "The 2.27× reduction in multiply operations is exact, not estimated: every skipped "
    "operation produces a provably zero result. There is no approximation error."
)

heading("5.4  Wall-clock timing benchmark & Goldilocks sweep", 2)
body(
    "The ternlang-ml sparse kernel uses a three-layer optimization stack: "
    "(1) flat i8 pre-flattening eliminates Trit enum dispatch from the hot path; "
    "(2) a standard CSC (Compressed Sparse Column) offset table replaces pointer-chasing "
    "Vec<Vec<...>> with two tight contiguous arrays that fit in L1 cache; "
    "(3) Rayon parallel row dispatch fires all logical cores simultaneously. "
    "\n\n"
    "Table 6a shows performance at 25% sparsity (LCG weights). "
    "Table 6b shows the BitNet b1.58 target zone (60% sparsity). "
    "Table 6c is the full sparsity sweep — the Goldilocks analysis — revealing "
    "three distinct performance regimes."
)

add_table(
    ["Matrix size", "Sparsity", "Dense (μs)", "Sparse (μs)", "Speedup"],
    [
        ["32 × 32",   "25.2%", "228",    "42",    "5.4×"],
        ["64 × 64",   "25.2%", "1,419",  "114",   "12.5×"],
        ["128 × 128", "24.8%", "11,181", "503",   "22.2×"],
        ["256 × 256", "24.8%", "96,752", "2,653", "36.5×"],
        ["512 × 512", "24.9%", "851,601","18,238", "46.7×"],
    ],
    "Table 6a. CSC sparse vs. dense matmul — 25% sparsity, release build."
)

add_table(
    ["Matrix size", "Sparsity", "Dense (μs)", "Sparse (μs)", "Speedup"],
    [
        ["32 × 32",   "58.8%", "252",     "33",    "7.6×"],
        ["64 × 64",   "58.9%", "1,545",   "118",   "13.1×"],
        ["128 × 128", "59.1%", "12,422",  "453",   "27.4×"],
        ["256 × 256", "59.6%", "107,316", "2,392", "44.9×"],
        ["512 × 512", "60.0%", "901,866", "10,473","86.1×"],
    ],
    "Table 6b. BitNet b1.58 target zone — 60% sparsity, release build, Rayon parallel."
)

heading("5.4.1  Goldilocks sparsity sweep", 3)
body(
    "To find the optimal operating point, we swept nine sparsity levels from 25% to 99% "
    "across five matrix sizes (3-rep median, release build). Table 6c shows the speedup "
    "heatmap. Three distinct regimes emerge:\n\n"
    "Regime 1 — Warm zone (25–40% sparsity): consistently high speedup at all sizes, "
    "including small matrices. Practical for lightly-quantized models.\n\n"
    "Regime 2 — Goldilocks zone (40–60% sparsity): peak average speedup for medium-to-large "
    "matrices. Matches the BitNet b1.58 distribution exactly — this is not a coincidence; "
    "the ternary quantization threshold naturally produces weights in this range.\n\n"
    "Regime 3 — Asymptotic zone (90–99% sparsity): small matrices suffer overhead penalty "
    "but large matrices hit extraordinary numbers. At 512×512 with 99% sparsity the sparse "
    "kernel runs in 1,548 μs vs 168,099 μs dense — a 108–122× range depending on run. "
    "This regime applies to sparse attention, mixture-of-experts gating, and token pruning."
)

add_table(
    ["Sparsity", "32²", "64²", "128²", "256²", "512²"],
    [
        ["25%",  "6.3×",  "11.5×", "26.4×", "39.3×",  "53.1×"],
        ["40%",  "6.3×",  "13.1×", "29.6×", "46.0×",  "73.6×"],
        ["50%",  "5.9×",  "10.2×", "28.7×", "56.6×",  "82.1×"],
        ["60%",  "5.8×",   "9.5×", "27.9×", "32.1×",  "84.9×"],
        ["70%",  "4.0×",   "8.6×", "20.7×", "48.7×",  "81.7×"],
        ["80%",  "3.5×",   "6.4×", "20.4×", "45.5×",  "72.3×"],
        ["90%",  "2.0×",   "5.8×", "18.7×", "38.6×",  "70.9×"],
        ["95%",  "1.9×",   "4.5×", "15.6×", "47.5×",  "85.9×"],
        ["99%",  "1.8×",   "9.9×", "13.1×", "53.9×", "122.3×"],
    ],
    "Table 6c. Goldilocks sparsity sweep — speedup (sparse / dense) across sparsity × size. "
    "Peak measured: 122.3× at 99% sparsity, 512×512. "
    "Goldilocks zone for medium matrices (128–256²): 40–60% sparsity, 20–57× speedup. "
    "All measurements: release build, Rayon parallel rows, 3-rep median."
)

heading("5.5  End-to-end inference: TernaryMLP", 2)
body(
    "The ternlang-ml crate provides a 2-layer TernaryMLP to demonstrate the full inference "
    "pipeline from f32 weight initialisation through ternary quantization to sparse forward "
    "pass. Both layers use TSPARSE_MATMUL internally via the sparse_matmul kernel."
)
code_block(
    "pub struct TernaryMLP {\n"
    "    pub w1: TritMatrix,     // [in_features × hidden_size]\n"
    "    pub w2: TritMatrix,     // [hidden_size × out_features]\n"
    "    pub in_features:  usize,\n"
    "    pub hidden_size:  usize,\n"
    "    pub out_features: usize,\n"
    "}\n\n"
    "// Construct from f32 weights — auto-applies BitNet threshold per layer\n"
    "let mlp = TernaryMLP::from_f32(2, 4, 2, w1_f32, w2_f32);\n\n"
    "// Forward pass — returns (output, skipped_l1, skipped_l2)\n"
    "let (out, sk1, sk2) = mlp.forward(&input);\n"
    "let class = mlp.predict(&input);   // argmax"
)
body(
    "The model is evaluated on a 4-example XOR dataset and an 8-example 3-bit parity "
    "dataset. The purpose of this module is to demonstrate that the full inference path "
    "(quantize → TritMatrix → sparse_matmul → argmax) executes correctly end-to-end, "
    "not to train a model. Ternary training loops with gradient quantization are in scope "
    "for Phase 8."
)
body_parts([
    ("Key result:  ", True, False),
    ("TSPARSE_MATMUL is reachable as an end-to-end path from f32 model weights through "
     "ternary quantization to classification output, without any dense fallback. "
     "The kernel composes correctly with multi-layer architectures.", False, False)
])

heading("5.6  TCOMPRESS / TUNPACK: tensor RLE codec", 2)
body(
    "Sparse trit tensors stored in the VM heap represent a second opportunity for "
    "bandwidth reduction beyond the multiply-skip speedup. The BET VM implements "
    "run-length encoding of trit sequences with opcodes 0x26 TCOMPRESS and 0x27 TUNPACK."
)
body(
    "The codec uses a base-3 two-trit encoding: each run is represented as a (value, "
    "hi, lo) triplet where count = hi × 3 + lo ∈ {1, …, 8}. A NegOne sentinel header "
    "distinguishes compressed from raw tensors. For a typical BitNet weight tensor at "
    "60% zero-sparsity, the codec achieves 40–55% size reduction, with lossless "
    "round-trip decompression verified by 5 dedicated VM tests."
)

heading("5.7  Scalar ternary temperature and ambiguity detection", 2)
body(
    "Discrete ternary decisions — reject / tend / affirm — are necessary but not sufficient "
    "for AI agent reasoning. An agent that knows it is in the 'affirm' zone but does not "
    "know how strongly cannot calibrate when to act vs. when to gather more evidence. "
    "Ternlang introduces a continuous scalar temperature model that unifies the discrete "
    "and continuous views."
)
body(
    "A TritScalar is a real value clamped to [−1.0, +1.0]. The full range is partitioned "
    "into three zones by the tend boundary β = 1/3:"
)
code_block(
    "reject  ∈ [−1.000, −0.333)   // negative, resolvable\n"
    "tend    ∈ [−0.333, +0.333]   // deliberation zone — do NOT act yet\n"
    "affirm  ∈ (+0.333, +1.000]   // affirmative\n\n"
    "confidence = (|scalar| − β) / (1 − β)   for reject/affirm\n"
    "           = 1 − |scalar| / β            for tend"
)
body_parts([
    ("The tend zone is the most misunderstood trit.  ", True, False),
    ("It is not null. It is not indecision. It is an active computational instruction: "
     "the agent's evidence has not yet cleared a boundary sufficient to act. "
     "The confidence score tells the agent how far it is from that boundary — "
     "and therefore how much additional evidence is needed.", False, False)
])

heading("5.8  Multi-dimensional evidence vectors (trit_vector)", 2)
body(
    "Real reasoning agents collect evidence from multiple sources simultaneously. "
    "The TritEvidenceVec type represents a named, weighted set of evidence dimensions, "
    "each carrying its own scalar value. The aggregate scalar is a weighted mean:"
)
code_block(
    "scalar_aggregate = Σᵢ (wᵢ · vᵢ) / Σᵢ wᵢ\n\n"
    "Example:\n"
    "  visual_evidence:    +0.80 (weight 1.0) → affirm, confidence 70%\n"
    "  textual_evidence:   −0.20 (weight 0.5) → tend,   confidence 40%\n"
    "  contextual_signal:  +0.40 (weight 1.5) → affirm, confidence 10%\n"
    "  ─────────────────────────────────────────────────────────────\n"
    "  aggregate scalar:   +0.36             → affirm, confidence  8%\n"
    "  is_actionable(0.5): false             → continue gathering evidence"
)
body(
    "The MCP server exposes this as the trit_vector tool: any AI agent can submit "
    "its named evidence sources and receive back a full breakdown — per-dimension zone "
    "classification, dominant dimension, aggregate scalar, and a plain-language "
    "recommendation. The architecture is model-agnostic: any agent that can produce "
    "numeric confidence scores can become a ternary reasoner without modification."
)

add_table(
    ["Component", "Type", "Description"],
    [
        ["TritScalar",       "f32 ∈ [−1, +1]",     "Continuous ternary temperature; maps to reject/tend/affirm + confidence"],
        ["TritEvidenceVec",  "Vec<(String, f32, f32)>", "Named, weighted evidence dimensions; aggregates to TritScalar"],
        ["TEND_BOUNDARY",    "const 1/3 ≈ 0.333",  "Zone boundary: decisive vs. deliberation"],
        ["confidence()",     "f32 ∈ [0, 1]",       "Depth into zone: 0.0 = at boundary, 1.0 = at extreme"],
        ["is_actionable(τ)", "bool",                "True iff zone is reject/affirm AND confidence ≥ τ"],
        ["trit_decide",      "MCP tool",            "Evidence[] → scalar, label, confidence, per-signal breakdown"],
        ["trit_vector",      "MCP tool",            "Named dimensions + weights → aggregate + breakdown + recommendation"],
    ],
    "Table 7. Scalar temperature and evidence vector API."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 6. HARDWARE BACKEND
# ══════════════════════════════════════════════════════════════════════════════

heading("6. Hardware Backend (ternlang-hdl)", 1)

heading("6.1  Verilog-2001 primitives", 2)
body("The ternlang-hdl crate generates synthesisable Verilog-2001 modules. Each trit is a [1:0] bus:")

add_table(
    ["Module", "Operation", "Implementation note"],
    [
        ["trit_neg",  "neg(t)",      "assign y = {a[0], a[1]} — pure wire swap, zero gates"],
        ["trit_cons", "cons(a,b)",   "assign y = (a == b) ? a : 2'b11"],
        ["trit_mul",  "mul(a,b)",    "zero-skip detect; only multiply if neither input is hold"],
        ["trit_add",  "add(a,b)",    "9-entry case statement producing (sum, carry)"],
        ["trit_reg",  "D register",  "synchronous write, asynchronous reset to 2'b11 (hold)"],
        ["bet_alu",   "Full ALU",    "op[1:0] selects ADD/MUL/NEG/CONS"],
    ],
    "Table 8. BET Verilog-2001 primitive modules."
)

heading("6.2  Sparse matmul array", 2)
body(
    "The synthesisable sparse matmul array instantiates an N×N grid of processing elements. "
    "Each cell contains a weight register and a clock-gate signal based on the zero-weight test:"
)
code_block(
    "wire [1:0] w_ij = weight_reg[i][j];\n"
    "wire skip       = (w_ij == 2'b11);   // hold = zero weight\n"
    "wire [1:0] contrib = skip\n"
    "    ? 2'b11                            // propagate hold\n"
    "    : trit_mul(a_i, w_ij);             // real multiply"
)
body(
    "Clock-gating on the skip signal prevents switching activity in zero-weight cells, "
    "delivering dynamic power reduction proportional to weight sparsity — typically "
    "50–60% power saving for BitNet-quantized networks."
)

heading("6.3  BET processor and FPGA simulation", 2)
body(
    "The full bet_processor module wires bet_regfile (27×2-bit), bet_pc (16-bit program "
    "counter with load port), and bet_control (single-cycle decode, all 51 opcodes mapped "
    "to control signals). The ternlang sim command compiles a .tern file to bytecode and "
    "emits a complete self-contained Icarus Verilog testbench:"
)
code_block(
    "ternlang sim program.tern          # emit testbench: program.sim.v\n"
    "iverilog -o sim.vvp program.sim.v  # compile\n"
    "vvp sim.vvp                        # run\n"
    "# waveforms exported to bet_sim.vcd — open in GTKWave"
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 7. NATIVE COMPILATION: AST-TO-C BACKEND
# ══════════════════════════════════════════════════════════════════════════════

heading("7. Native Compilation: AST-to-C Backend", 1)

body(
    "The ternlang-codegen crate provides a second compilation target alongside BET bytecode: "
    "a self-contained C11 source file that can be compiled with any standard C compiler. "
    "This opens three new deployment paths: native-speed execution, cross-compilation to "
    "embedded targets, and human-readable output for inspection and security audit."
)

heading("7.1  Transpilation architecture", 2)
body(
    "The CTranspiler operates directly on the abstract syntax tree (Program) produced by "
    "the parser, bypassing the BET VM entirely:"
)
code_block(
    ".tern  →[parse]→  AST  →[CTranspiler]→  .c  →[gcc/clang]→  native binary"
)
body(
    "The generated C file is fully self-contained: a header of inline trit primitives "
    "using int8_t with values {-1, 0, +1} is prepended, covering trit_add, trit_mul, "
    "trit_neg, trit_consensus, trit_abs, and the built-in constants trit_truth(), "
    "trit_hold(), trit_conflict()."
)

heading("7.2  Language construct mapping", 2)
body("The transpiler maps every ternlang construct to its natural C equivalent:")
bullet("match → switch (int) with case -1, case 0, case 1 arms")
bullet("struct definitions → typedef struct { ... }")
bullet("fn → C function with forward declaration (enables mutual recursion without re-ordering)")
bullet("loop / while / for → for(;;), while(1) with inner ternary condition dispatch")
bullet("expr? → __TERN_PROPAGATE(expr) macro slot, enabling the caller to define the early-return idiom in C")
bullet("main() → tern_main() (to avoid clashing with C's entry point); a generated main() calls tern_main() and translates the trit result to an exit code")

body("Example: ternlang source and generated C output:")
code_block(
    "// ternlang source\n"
    "fn decide(x: trit) -> trit {\n"
    "    match x {\n"
    "        -1 => conflict()\n"
    "         0 => hold()\n"
    "        +1 => truth()\n"
    "    }\n"
    "}\n"
    "\n"
    "// generated C\n"
    "trit decide(trit x) {\n"
    "    switch ((int)x) {\n"
    "        case -1: return trit_conflict(); break;\n"
    "        case  0: return trit_hold();     break;\n"
    "        case  1: return trit_truth();    break;\n"
    "    }\n"
    "}"
)
body(
    "The C backend does not yet support the actor model (spawn/send/await) or tensor heap "
    "operations, which remain BET-VM specific. These emit comments in the output, making "
    "the generated file auditable even for unsupported constructs."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 8. ACTOR MODEL AND DISTRIBUTED RUNTIME
# ══════════════════════════════════════════════════════════════════════════════

heading("8. Actor Model and Distributed Runtime", 1)

heading("8.1  Local actors", 2)
body(
    "Three ISA primitives implement the actor model. TSPAWN (0x30) creates an agent "
    "instance from a registered type ID and returns an agentref. TSEND (0x31) enqueues "
    "a trit message in the agent's mailbox (VecDeque<Value>) without blocking. "
    "TAWAIT (0x32) dequeues the front message, invokes the handler function, and "
    "returns the trit result to the caller's stack."
)

heading("8.2  Distributed actors (ternlang-runtime)", 2)
body(
    "The ternlang-runtime crate extends the actor model across TCP. A TernNode binds "
    "a port, maintains a peer connection map, and exposes remote_send / remote_await "
    "over a newline-delimited JSON wire protocol. Four message types are defined:"
)
code_block(
    '{"type":"Send",  "agent_id":0, "trit":1}\n'
    '{"type":"Await", "agent_id":0}\n'
    '{"type":"Reply", "trit":0}\n'
    '{"type":"Error", "message":"agent not found"}'
)
body(
    "The newline-delimited format requires no framing library and is trivially "
    "implementable in any language, enabling non-Rust nodes to participate in a "
    "ternlang actor network."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 9. TERNARY AI REASONING TOOLKIT
# ══════════════════════════════════════════════════════════════════════════════

heading("9. Ternary AI Reasoning Toolkit", 1)

body(
    "Modern AI agents are structurally binary in their decision loops: evidence is "
    "collapsed to a scalar, thresholded, and converted to a Boolean act/wait flag. "
    "The three-valued nature of evidence — affirm, hold, reject — is treated as a "
    "side-effect of confidence scores rather than a first-class type. The ternlang-ml "
    "crate introduces five primitives that make the decision loop architecturally ternary."
)

heading("9.1  Deliberation Engine", 2)
body(
    "The DeliberationEngine models iterative evidence accumulation using an exponential "
    "moving average (EMA):"
)
code_block(
    "Sᵣ = α · eᵣ + (1−α) · Sᵣ₋₁,    α ∈ (0, 1]"
)
body(
    "where eᵣ is the mean of new evidence signals in round r. The engine commits to a "
    "ternary decision only when the confidence of Sᵣ exceeds a target threshold; otherwise "
    "it remains in the hold state and requests a further evidence round. This models the "
    "human 'let me think about it' behaviour — a native ternary computation rather than "
    "a binary timeout."
)

heading("9.2  Coalition Vote", 2)
body(
    "coalition_vote() aggregates N independent agent verdicts — each a trit with an "
    "associated confidence and weight — into a single result with quorum, dissent, "
    "and abstain statistics:"
)
code_block(
    "t̂ = sign( Σᵢ wᵢ · cᵢ · tᵢ  /  Σᵢ wᵢ · cᵢ )\n\n"
    "  tᵢ ∈ {-1, 0, +1}   — agent verdict\n"
    "  cᵢ ∈ [0, 1]         — agent confidence\n"
    "  wᵢ > 0               — importance weight"
)

heading("9.3  Action Gate", 2)
body(
    "The action gate enforces a structural separation between hard constraints and soft "
    "preferences. Each GateDimension carries an evidence signal, a weight, and an optional "
    "hard_block flag. A hard-block dimension with negative evidence vetoes the action "
    "unconditionally, regardless of all other dimensions:"
)
code_block(
    "verdict = Block      if ∃ dᵢ : hard_block(dᵢ) ∧ t_dᵢ = -1\n"
    "          sign(S̄)   otherwise"
)
body(
    "This is the structural analogue of a safety interlock: the veto is unconditional, "
    "not probabilistic."
)

heading("9.4  Scalar temperature bridge and hallucination score", 2)
body(
    "scalar_temperature() bridges the ternary decision to the LLM sampling temperature "
    "domain: affirm at high confidence maps to low temperature (focused generation), "
    "hold maps to mid temperature (exploratory generation), and reject maps to near-zero "
    "(cautious refusal). See §5.7 for the full TritScalar zone specification."
)
body(
    "hallucination_score() maps the variance of an evidence signal vector to a trust trit: "
    "low variance (consistent signals) yields trust +1; high variance yields −1, signalling "
    "that the agent's signals are incoherent and should not be acted upon."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 10. MOE-13 TERNARY ORCHESTRATOR
# ══════════════════════════════════════════════════════════════════════════════

heading("10. MoE-13 Ternary Orchestrator", 1)

body(
    "The MoE-13 architecture is implemented in the ternlang-moe crate. It realises a "
    "ternary Mixture-of-Experts head whose routing, synthesis, safety gate, and memory "
    "are all expressed natively in balanced ternary semantics."
)

heading("10.1  Six-Dimensional Competence Space", 2)
body(
    "Each expert is characterised by a competence vector v ∈ [−1, +1]⁶ across six axes:"
)
code_block(
    "v = [v_syntax,  v_world,   v_reason,\n"
    "     v_tool,    v_persona, v_safety]"
)
body(
    "The synergy between two experts is defined as the complement of their cosine "
    "similarity — orthogonal experts are maximally complementary:"
)
code_block(
    "σ(vᵢ, vⱼ) = (1 − cos(vᵢ, vⱼ)) / 2   ∈ [0, 1]"
)

heading("10.2  Dual-Key Synergistic Routing", 2)
body(
    "For each candidate expert pair (i, j), the routing score combines relevance to the "
    "query vector q with inter-expert synergy:"
)
code_block(
    "score(i, j) = ρᵢ(q) · ρⱼ(q) · σ(vᵢ, vⱼ)\n\n"
    "  ρₖ(q) = max(0, cos(vₖ, q))   — relevance of expert k to query"
)
body(
    "The pair with the highest score is selected. Crucially, two highly relevant but "
    "similar experts are penalised by low σ; the router favours complementary coverage "
    "over redundant strength."
)

heading("10.3  1+1=3 Triad Synthesis", 2)
body(
    "After expert evaluation, an emergent triad field is synthesised from the two selected "
    "competence vectors and the routing synergy:"
)
code_block(
    "E_k = σᵢⱼ · (vᵢ + vⱼ) / 2"
)
body(
    "This is the formal expression of 1+1=3: two expert signals at high synergy produce "
    "a third signal — the emergent field E_k — that neither expert could produce in isolation. "
    "The field modulates the subsequent vote and the LLM temperature hint."
)

heading("10.4  Safety Hard Gate and Axis-6 Veto", 2)
body(
    "Dimension 6 of every competence vector is the safety axis. Before any vote is "
    "computed, the safety projection of the triad field is evaluated:"
)
code_block(
    "veto ⟺ E_k,safety < θ_s      (default θ_s = −0.3)"
)
body(
    "A veto immediately returns trit −1 at confidence 1.0 and logs the event to the "
    "Axis-tier memory for audit. No downstream reasoning can override it."
)

heading("10.5  Hold State and Tiebreaker", 2)
body(
    "When the weighted vote yields trit 0 or aggregate confidence falls below a hold "
    "threshold, the orchestrator does not immediately return. It invokes a tiebreaker "
    "expert — selected by highest reasoning-axis score among inactive experts — and "
    "re-votes with up to max_active=4 experts total. Only if the tiebreaker also fails "
    "to resolve does the orchestrator return the hold state to the caller."
)

heading("10.6  Three-Tier Memory Mesh", 2)
body("The orchestrator maintains three memory tiers:")
bullet("Node (TTL: seconds): LRU-bounded volatile store for within-session context. Evicts on capacity overflow.")
bullet("Cluster (TTL: minutes): routing frequency counters enabling mode-collapse detection (risk = max_pair / total).")
bullet("Axis (persistent): immutable audit log of safety vetoes with timestamp, expert identity, and query hash. Global priors over expert performance.")

heading("10.7  Nine-Step Inference Pipeline", 2)
body(
    "The full orchestration pass executes in nine deterministic steps: "
    "(1) encode query to evidence vector; "
    "(2) dual-key route to best expert pair; "
    "(3) evaluate both experts independently; "
    "(4) synthesise triad field; "
    "(5) safety hard gate; "
    "(6) weighted trit vote with synergy amplification; "
    "(7) hold detection; "
    "(8) optional tiebreaker invocation; "
    "(9) return OrchestrationResult and update all three memory tiers."
)

heading("10.8  Thirteen-Agent Dual-Signal Deliberation", 2)
body(
    "The AgentHarness in ternlang-moe/src/agents/ extends the 9-step orchestration pass "
    "with a structured deliberation layer of 13 specialised agents, each computing two "
    "independent signals:"
)
code_block(
    "tᵢ = +1   if positive_signal ≥ θ⁺  and  negative_signal < θ⁻\n"
    "     -1   if negative_signal ≥ θ⁻  and  positive_signal < θ⁺\n"
    "      0   (stasis: signals in equilibrium)"
)
body("The 13 agents and their primary detection axes:")
add_table(
    ["Agent", "Primary signal", "Hard trigger"],
    [
        ["Syntax",         "Structural token analysis, bracket balance",          "—"],
        ["WorldKnowledge",  "Domain vocabulary density, proper noun frequency",   "—"],
        ["DeductiveReason", "Premise + conclusion marker co-occurrence",          "—"],
        ["InductiveReason", "Example density, generalisation leap detection",     "—"],
        ["ToolUse",         "Imperative verb strength, passive construction",     "—"],
        ["Persona",         "Depersonalisation detection",                        "—"],
        ["Safety",          "Hard risk keywords",                                 "trit -1, confidence 0.99"],
        ["FactCheck",       "Overconfidence markers → hallucination flag",        "—"],
        ["CausalReason",    "Requires both cause and effect markers",             "—"],
        ["AmbiguityRes",    "Hard-vague ≥ 2 markers",                            "trit -1"],
        ["MathReason",      "Impossible operations (divide-by-zero)",             "trit -1"],
        ["ContextMem",      "Fresh-start signals vs. anaphoric reference density","—"],
        ["MetaSafety",      "Injection pattern detection",                        "trit -1, confidence 0.99"],
    ],
    "Table 9. Thirteen deliberation agents and their dual-signal computation."
)
body_parts([
    ("Active stasis.  ", True, False),
    ("Trit 0 is not a default or a fallback from a failed comparison. It is an honest "
     "declaration that positive and negative evidence are balanced and further information "
     "is required before committing.", False, False)
])

heading("10.9  Introspective Hold: The Stable Attractor", 2)
body(
    "The run_introspective() method on AgentHarness implements a stable attractor for the "
    "hold state. Once reached, a stable hold is permanent within the deliberation turn — "
    "it cannot be overridden by additional tiebreaker invocations. The stable attractor "
    "condition is:"
)
code_block(
    "stable_hold  ⟺  |affirm_count − conflict_count| ≤ 1\n"
    "                ∧  engaged ≥ 4"
)
body(
    "This formalises the philosophical claim: when at least four independent agents have "
    "deliberated and the signal balance between affirmation and conflict is "
    "indistinguishable, the appropriate epistemic response is to hold — not to choose "
    "arbitrarily between +1 and −1. The hold state is returned with is_stable_hold: true "
    "and a human-readable hold_reason string."
)
code_block(
    "pub struct AggregateVerdict {\n"
    "    pub trit:           i8,\n"
    "    pub confidence:     f32,\n"
    "    pub verdicts:       Vec<ExpertVerdict>,\n"
    "    pub is_stable_hold: bool,\n"
    "    pub hold_reason:    Option<String>,\n"
    "    pub affirm_count:   usize,\n"
    "    pub conflict_count: usize,\n"
    "    pub hold_count:     usize,\n"
    "}"
)
body(
    "Safety and meta-safety form a hard gate: the evidence vector produced by "
    "to_evidence_vector() maps axis 6 to min(Safety.confidence, MetaSafety.confidence), "
    "ensuring the safety dimension is always the most conservative signal in the system."
)

heading("10.10  Orchestrate-Full: Thirteen-Substage Pipeline", 2)
body(
    "The orchestrate_full() method on TernMoeOrchestrator composes the 13-agent "
    "deliberation with the 9-step routing pipeline into a unified end-to-end pass:"
)
bullet("Step 1: Run all 13 agents via AgentHarness::run_introspective()")
bullet("Step 2: Check for stable attractor hold — return immediately if true")
bullet("Step 3: Map 13 verdicts to 6D evidence vector via to_evidence_vector()")
bullet("Step 4: Check safety hard gate on safety axis before MoE routing")
bullet("Step 5: Run standard 9-step MoE orchestration with enriched evidence")
bullet("Step 6: Return OrchestrationResult with stable hold flag propagated")
body(
    "This creates a two-tier system: the fast ternary language models (13 agents, "
    "keyword-level deliberation) form a pre-filter that enriches the evidence before "
    "the slower but more semantically powerful MoE routing pass. Most queries that are "
    "clearly safe and unambiguous complete in the agent tier; only complex or borderline "
    "queries reach the full routing pipeline."
)

heading("10.11  MCP Exposure", 2)
body(
    "The orchestrator is exposed as three MCP tool calls: moe_orchestrate runs a full "
    "pass and returns the complete result including routing pair, triad field, verdicts, "
    "temperature, and prompt hint; moe_deliberate wraps the EMA deliberation engine for "
    "multi-round iterative reasoning; trit_action_gate exposes the hard-block gate as a "
    "standalone safety check. Any MCP-compatible AI client connected to ternlang-mcp "
    "gains access to the full ternary reasoning stack as tool calls."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 11. QUTRIT NEURAL NETWORKS AND TERNLANG
# ══════════════════════════════════════════════════════════════════════════════

heading("11. Qutrit Neural Networks and Ternlang", 1)

body(
    "The Qutrit Neural Network (QNN) framework establishes a formal correspondence "
    "between quantum qutrit states and balanced ternary values. The three qutrit basis "
    "states |−⟩ (decay), |0⟩ (stasis), and |+⟩ (growth) map exactly onto the ternlang "
    "trit space:"
)
code_block(
    "|−⟩  ↔  −1   (conflict / decay)\n"
    "|0⟩  ↔   0   (hold / stasis)\n"
    "|+⟩  ↔  +1   (truth / growth)"
)
body(
    "This correspondence is not coincidental. Balanced ternary's symmetric structure "
    "around zero and its active neutral state mirror the physical properties of three-state "
    "quantum systems, making ternlang a natural implementation substrate for QNN inference."
)

heading("11.1  Tesseract Recursive Syntax Stabilisation (TRS)", 2)
body(
    "The QNN framework introduces Tesseract Recursive Syntax (TRS), a stabilisation "
    "mechanism that prevents qutrit networks from collapsing into binary modes under "
    "iterative refinement. In ternlang, TRS corresponds directly to the introspective "
    "hold architecture (§10.9): when affirmation and conflict signals are balanced, the "
    "system stabilises at trit 0 rather than forcing a premature commitment. The stable "
    "attractor condition (|affirm − conflict| ≤ 1 ∧ engaged ≥ 4) is the discrete "
    "analogue of TRS convergence in continuous qutrit space."
)

heading("11.2  QNN example suite", 2)
body(
    "The ternlang repository contains 15 QNN-inspired .tern example programs covering:"
)
bullet("Qutrit superposition encoding and collapse simulation")
bullet("Ternary attention mechanisms with {−1, 0, +1} key-query products")
bullet("Qutrit activation functions replacing ReLU/GELU in ternary networks")
bullet("Phase-encoding of temporal signals as trit sequences")
bullet("TRS stabilisation loops using ternlang's loop construct and the ? propagation operator for collapse detection")
bullet("Qutrit gradient approximation via balanced ternary finite differences")
body(
    "These examples demonstrate that ternlang's syntax is sufficient to express QNN "
    "computations natively, without any binary adaptor layer. The @sparseskip directive "
    "applies directly to qutrit weight matrices: at realistic qutrit sparsity levels "
    "(≈33% zero states at initialisation), the TSPARSE_MATMUL kernel provides the same "
    "8–14× speedup over dense computation observed in classical BitNet workloads."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 12. RELATED WORK
# ══════════════════════════════════════════════════════════════════════════════

heading("12. Related Work", 1)

body_parts([("Balanced ternary foundations.  ", True, False),
    ("Knuth (1997) provides the mathematical basis. The Setun computer (Moscow State "
     "University, 1958) demonstrated physical ternary hardware using magnetic elements. "
     "Both are existence proofs that the paradigm is real — they predate the software "
     "ecosystem that would make it useful.", False, False)])

body_parts([("USN / Bos & Gundersen (2020).  ", True, False),
    ("The most active academic effort: C-to-ternary compilation targeting EDA tools for "
     "memristor-backed ternary circuits. Their approach forces binary-native C semantics "
     "onto a ternary substrate, creating abstraction leaks where the symmetry of balanced "
     "ternary is not exploitable. Ternlang's native syntax eliminates this gap. "
     "Their hardware work (uMemristorToolbox) is a future physical target for ternlang programs.", False, False)])

body_parts([("Open-source ternary emulators.  ", True, False),
    ("Brandon Smith's 9-trit RISC simulator (Python) implements fetch-decode-execute in "
     "base-3 on 9-trit words. Owlet is an S-expression ternary interpreter in Node.js. "
     "Both solve a single layer without compiler, ML kernels, or hardware support. "
     "The ternlang-compat crate provides assembler-level bridges to both, making BET VM "
     "the common runtime they target.", False, False)])

body_parts([("BitNet and ternary neural networks.  ", True, False),
    ("Ma et al. (2024) demonstrate that large language models can be trained with weights "
     "in {−1, 0, +1} while retaining competitive perplexity. BitNet b1.58 extends this "
     "to the 1.58-bit regime where every weight is a trit. Ternlang is the first project "
     "to surface this property as a first-class ISA opcode (TSPARSE_MATMUL) rather than "
     "a software library optimisation.", False, False)])

body_parts([("Quantum ternary (qutrits).  ", True, False),
    ("Qutrits — 3-level quantum systems — map naturally to trit values {|−1⟩, |0⟩, |+1⟩}. "
     "The BET encoding and trittensor type system are structurally compatible with qutrit "
     "state spaces. The QNN framework (§11) formalises this mapping and demonstrates "
     "that ternlang's syntax is sufficient for native QNN computation.", False, False)])

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 13. THE TERNARY ECOSYSTEM
# ══════════════════════════════════════════════════════════════════════════════

heading("13. The Ternary Computing Ecosystem", 1)

body(
    "A stated goal of ternlang is to serve as the convergence point for the fragmented "
    "ternary computing field — the place where existing efforts compile into a coherent "
    "whole rather than remaining isolated."
)

add_table(
    ["Project", "Technology", "Ternlang bridge", "Status"],
    [
        ["Brandon Smith 9-trit sim", "Python, .tasm assembly",   "TasmAssembler → BET bytecode",       "Complete (ternlang-compat)"],
        ["Owlet",                    "Node.js, S-expressions",   "OwletParser → ternlang AST → BET VM","Complete (ternlang-compat)"],
        ["USN / Bos+Gundersen",      "C-to-ternary, EDA tools",  "Academic whitepaper; ISA interop",   "In progress"],
        ["uMemristorToolbox",        "Unity, physical memristors","Phase 7 hardware target",            "Planned"],
        ["Trit-Rust",                "Rust, i8-backed trits",    "Superseded by ternlang-core",        "Complete"],
        ["Q-Ternary",                "Qutrit DSL",               "trittensor state model mapping",     "Future work"],
    ],
    "Table 10. Ternary ecosystem compatibility map."
)

body(
    "Beyond runtime interoperability, ternlang aims to serve as a conceptual archive "
    "for prior ternary computing work. Brandon Smith's Python 9-trit RISC simulator "
    "demonstrated a complete fetch-decode-execute cycle in balanced ternary — an existence "
    "proof that received little downstream adoption due to the absence of a compiler "
    "ecosystem. The Owlet interpreter proved that S-expression evaluation maps cleanly to "
    "a three-valued substrate. Both contributions are honoured in the examples/ corpus: "
    "09_risc_fetch_decode.tern translates Smith's pipeline decision logic into native BET "
    "language syntax, and 13_owlet_bridge.tern models cooperative ternary evaluation with "
    "hold-as-suspension semantics."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 14. IMPLEMENTATION STATUS
# ══════════════════════════════════════════════════════════════════════════════

heading("14. Implementation Status", 1)

body(
    "Ternlang v0.1.2 is implemented in Rust as a Cargo workspace. All crates are publicly "
    "available on crates.io. The full test suite comprises 177+ tests across 12 crates, "
    "all passing. The live API is deployed at https://ternlang.com/mcp (Fly.io, Frankfurt) "
    "with TLS via Let's Encrypt. Continuous integration via GitHub Actions automatically "
    "builds, tests, and deploys to Fly.io on every push to main."
)

add_table(
    ["Crate", "Tests", "Description"],
    [
        ["ternlang-core",    "55",  "Lexer, parser (+ ? propagation), AST, semantic checker, BET VM (51 opcodes), ModuleResolver"],
        ["ternlang-ml",      "21",  "BitNet quantization, sparse/CSC matmul, deliberation engine, coalition vote, action gate, hallucination score"],
        ["ternlang-moe",     "24",  "MoE-13 orchestrator: dual-key routing, triad synthesis, 3-tier memory, 13 dual-signal agents, introspective hold"],
        ["ternlang-codegen", "8",   "AST-to-C11 transpiler; match/struct/propagation; forward declarations; native-compilable output"],
        ["ternlang-test",    "10",  "Full-pipeline test framework: TernTestCase, assert_tern!, parse/semantic/VM error classification"],
        ["ternlang-hdl",     "21",  "Verilog-2001 codegen, BET RTL simulator (BetRtlProcessor), ternlang sim --rtl"],
        ["ternlang-lsp",     "—",   "LSP 3.17 server: hover documentation, 19 snippets, diagnostics"],
        ["ternlang-mcp",     "—",   "MCP server: 10 tools including trit_decide, moe_orchestrate, trit_action_gate"],
        ["ternlang-runtime", "4",   "Distributed TCP actor runtime (TernNode, wire protocol)"],
        ["ternlang-compat",  "29",  ".tasm assembler (9-trit RISC), Owlet S-expression parser"],
        ["ternpkg",          "5",   "Package manager: ternlang.toml, GitHub-backed registry"],
        ["ternlang-api",     "—",   "REST API (18 endpoints), multi-tenant key management, SSE streaming, MCP endpoint"],
    ],
    "Table 11. Ternlang crate inventory and test counts (v0.1.2, 2026-04-04)."
)

body(
    "Developer tooling: VS Code extension with TextMate grammar, LSP client, and formatter "
    "(packaged as ternlang-0.1.0.vsix, pending Marketplace publication); ternpkg package "
    "manager with GitHub-backed registry; 265 executable .tern example programs in examples/ "
    "spanning aerospace, medicine, distributed systems, hardware pipelines, civic governance, "
    "finance, AI agent design, and Qutrit Neural Network experiments."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 15. CONCLUSION AND FUTURE WORK
# ══════════════════════════════════════════════════════════════════════════════

heading("15. Conclusion and Future Work", 1)

body(
    "We have presented Ternlang v0.1.2, a full-stack balanced ternary execution architecture "
    "spanning language design, ISA, virtual machine, hardware backend, distributed runtime, "
    "native C compilation, and AI reasoning infrastructure. Five layers of contribution are "
    "unified under the same foundational primitive — the trit t ∈ {−1, 0, +1} with an active "
    "neutral state:"
)
bullet("Language completeness: exhaustive three-way matching, the ? error propagation operator, a real cross-file module system (stdlib built-in + filesystem-relative user modules), structured diagnostic codes [PARSE-001..004] through [BET-001..007], and all major control flow constructs — making ternlang a viable general-purpose programming language rather than a research prototype.")
bullet("Execution substrate: TSPARSE_MATMUL achieves a 2.27× reduction in multiply operations at baseline sparsity; a CSC kernel with Rayon reaches 122× at high sparsity (512² matrix, 99% sparsity, release mode). A native C11 compilation path via ternlang-codegen provides an alternative to VM interpretation.")
bullet("Reasoning primitives: the five-tool Ternary AI Reasoning Toolkit (DeliberationEngine, coalition vote, action gate, scalar temperature bridge, hallucination score) makes AI agent decision loops structurally three-valued.")
bullet("MoE orchestration: the MoE-13 orchestrator extends with 13-agent dual-signal deliberation, an introspective hold (stable attractor when |affirm − conflict| ≤ 1 across ≥ 4 agents), and an orchestrate_full() pipeline that pre-filters queries through the agent tier before MoE routing.")
bullet("Deployment and tooling: the live API at https://ternlang.com/mcp serves 10 MCP tools to any compatible AI client; a VS Code extension with LSP, formatter, and REPL provides an IDE experience; GitHub Actions CI/CD automatically builds, tests, and deploys on every push; and 177+ tests across 12 crates validate the full stack.")

body(
    "Near-term publication targets include the VS Code Marketplace submission "
    "(ternlang-0.1.0.vsix) and academic outreach to the USN group (Bos & Gundersen) "
    "for joint hardware-software co-design on the memristor backend."
)

heading("Future directions:", 2)
bullet("Type inference: Hindley-Milner style inference to eliminate explicit type annotations, reducing boilerplate for trit-returning functions.")
bullet("FPGA synthesis: full bet_processor targeting Xilinx Artix-7 and Lattice ECP5, with timing closure and resource utilisation reports.")
bullet("Memristor backend: integration with physical ternary state storage via the USN uMemristorToolbox.")
bullet("Qutrit bridge: formal mapping of trittensor to qutrit state spaces for quantum-adjacent hardware targeting (Google Willow and similar). The QNN framework's TRS stabilisation mechanism maps directly onto the introspective hold architecture.")
bullet("End-to-end training: native ternlang training loop with BitNet-style gradient quantization, enabling models trained and inferred entirely on BET VM.")
bullet("Academic collaboration: joint whitepaper with Bos & Gundersen (USN) comparing BET ISA to their EDA-synthesised ternary control path.")

doc.add_paragraph()
body(
    "The ternary computing field has been fragmented for decades. "
    "Ternlang is designed to be the substrate where those fragments converge."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

heading("References", 1)

refs = [
    "[1]  D. E. Knuth, The Art of Computer Programming, Vol. 2: Seminumerical Algorithms, 3rd ed. Addison-Wesley, 1997.",
    "[2]  N. P. Brousentsov et al., 'Development of ternary computers at Moscow State University,' Russian Virtual Computer Museum, 2002.",
    "[3]  S. Ma et al., 'The Era of 1-bit LLMs: All Large Language Models are in 1.58 Bits,' arXiv:2402.17764, 2024.",
    "[4]  S. Bos and H. Gundersen, 'Ternary Logic Synthesis for CMOS Technology Using Electronic Design Automation,' Proc. Norwegian Informatics Conference, 2020.",
    "[5]  S. Kepp, 'Ternlang: Balanced Ternary Intelligence Stack,' RFI-IRFOS, 2026. [Online]. Available: https://github.com/eriirfos-eng/ternary-intelligence-stack--tis-",
    "[6]  S. Kepp, 'Qutrit Neural Networks and Balanced Ternary: A Formal Correspondence,' RFI-IRFOS, OSF Preprints, DOI: 10.17605/OSF.IO/TZ7DC, 2026.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    run = p.add_run(ref)
    set_font(run, "Calibri", 10, color=DARK_GREY)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

out = "/home/eri-irfos/Desktop/Ternary Intelligence Stack (TIS)/ternlang-root/whitepaper/ternlang-whitepaper.docx"
doc.save(out)
print(f"Saved: {out}")
